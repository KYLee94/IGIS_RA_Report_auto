from __future__ import annotations

import hashlib
import json
import math
import re
import shutil
import zipfile
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter


ROOT = Path(__file__).resolve().parents[1]
IOTA_DIR = ROOT / "iota_reference"
FINANCE_CORE_DIR = IOTA_DIR / "finance" / "핵심"

MASTER_V09 = IOTA_DIR / "IOTA_Seoul_Master_DB_v0.9.xlsx"
UI_V09 = IOTA_DIR / "IOTA_Seoul_DB_for_UI_v0.9.xlsx"
HANDOFF_V09 = IOTA_DIR / "IOTA_Seoul_IFPDP_HandOff_v0.9.xlsx"

MASTER_V010 = IOTA_DIR / "IOTA_Seoul_Master_DB_v0.10.xlsx"
UI_V010 = IOTA_DIR / "IOTA_Seoul_DB_for_UI_v0.10.xlsx"
HANDOFF_V010 = IOTA_DIR / "IOTA_Seoul_IFPDP_HandOff_v0.10.xlsx"
REPORT_V010 = IOTA_DIR / "v0.10_build_report.md"

PUBLIC_HTML = ROOT / "public-pages" / "iota-development-workspace.html"
DOCS_HTML = ROOT / "docs" / "iota-development-workspace.html"
DATA_DIR = ROOT / "docs" / "data"
PACKAGE_DIR = ROOT / "docs" / "iota-data-package"

EXTRA_SHEET = "98_EXTRA_SOURCE_EXTRACTS_v0.10"
SOURCE_INDEX_SHEET = "v0.10_Source_File_Index"
NUMERIC_LOG_SHEET = "v0.10_Numeric_Candidate_Log"
LATESTNESS_LOG_SHEET = "v0.10_Latestness_Decision_Log"
LFC_DATA_SHEET = "v0.10_LFC_DATA"
UNFILLED_SHEET = "99_UNFILLED_FIELDS"

EXTRA_HEADERS = [
    "Candidate_ID",
    "Source_ID",
    "Source_Date",
    "Document_Version",
    "대분류",
    "중분류",
    "소분류",
    "항목",
    "값",
    "단위",
    "근거",
    "원문 존재 확실성",
    "사업 확정 상태",
    "Target_Sheet",
    "Target_PK",
    "Scenario_Tag",
    "Value_Status",
    "Promotion_Status",
    "Conflict_Group",
    "Superseded_By",
    "Extraction_Quality",
]

HANDOFF_HEADERS = ["대분류", "중분류", "소분류", "항목", "값", "단위", "상태", "화면추천위치"]


REFI_816_ROWS = [
    ("Tr.A-1", "케이에이치엘제이십일차㈜ (메리츠증권 SPC)", 1800),
    ("Tr.A-1", "메리츠화재", 1800),
    ("Tr.A-2", "갤럭시이오㈜ (NH투자증권 SPC)", 1300),
    ("Tr.B", "816공간제일차(신한증권)", 50),
    ("Tr.B", "한화실버아이언제일차(한화저축은행)", 50),
    ("Tr.B", "한국투자Debt Strategy 일반사모부동산투자신탁1호 (한투리얼에셋운용)", 600),
    ("Tr.B", "한국투자메자닌일반사모부동산투자신탁2호 (한투리얼에셋운용)", 350),
    ("Tr.B", "스틱얼터너티브자산운용㈜ (스틱크레딧안정화일반사모투자신탁제5호(전문))", 100),
    ("Tr.B", "대신저축은행", 80),
    ("Tr.B", "비씨카드", 150),
    ("Tr.B", "흥국저축은행", 20),
    ("Tr.C", "816공간제일차(신한증권)", 200),
    ("Tr.C", "키움가치추구형일반사모부동산투자신탁제1호 (키움투자자산운용)", 90),
    ("Tr.C", "이터널하이브(대신증권)", 480),
    ("Tr.C", "코람코국내개발일반사모부동산투자신탁제1-2호 (코람코운용)", 200),
    ("Tr.D", "케이에이치엘제이십이차㈜ (소노인터네셔널 SPC)", 700),
]

EQUITY_ROWS = [
    ("427", "보통주/우선주", "이지스REF", 591, "기투입"),
    ("427", "보통주/우선주", "이지스자산운용", 10, "기투입"),
    ("427", "보통주/우선주", "신한GIB", 100, "기투입"),
    ("427", "보통주", "신한투자증권", 30, "기투입"),
    ("427", "보통주", "신한은행", 50, "기투입"),
    ("427", "보통주", "신한캐피탈", 20, "기투입"),
    ("427", "보통주", "현대건설", 99, "기투입"),
    ("816", "보통주/우선주", "이지스421호", 19.55, "기투입"),
    ("816", "보통주", "이지스자산운용", 0, "기투입"),
    ("816", "보통주/우선주", "신한투자증권", 5.45, "기투입"),
    ("816", "보통주", "에셀유한회사", 0, "기투입"),
    ("816", "보통주", "NH투자증권", 0, "기투입"),
    ("816", "보통주", "삼성물산", 0, "기투입"),
]

EQUITY_427_DETAIL_ROWS = {"신한투자증권", "신한은행", "신한캐피탈"}


def find_file(name_part: str) -> Path | None:
    key = name_part.lower()
    matches = [p for p in IOTA_DIR.rglob("*") if p.is_file() and key in p.name.lower()]
    return sorted(matches, key=lambda p: (p.stat().st_mtime, len(str(p))), reverse=True)[0] if matches else None


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def infer_source_date(text: str, mtime: float | None = None) -> str:
    candidates = []
    for m in re.finditer(r"(20\d{6})", text):
        try:
            candidates.append(datetime.strptime(m.group(1), "%Y%m%d").date().isoformat())
        except ValueError:
            pass
    for m in re.finditer(r"(?<!\d)(\d{6})(?!\d)", text):
        try:
            candidates.append(datetime.strptime(m.group(1), "%y%m%d").date().isoformat())
        except ValueError:
            pass
    if candidates:
        return max(candidates)
    if "v2.0" in text or "PF Teaser" in text:
        return "2026-04-27"
    if mtime:
        return datetime.fromtimestamp(mtime).date().isoformat()
    return ""


def infer_version(text: str) -> str:
    m = re.search(r"(v\d+(?:\.\d+)?)", text, flags=re.IGNORECASE)
    if m:
        return m.group(1)
    if re.search(r"(?<![A-Za-z0-9])F2(?![A-Za-z0-9])", text):
        return "F2"
    return ""


def source_priority(path: Path) -> int:
    name = path.name.lower()
    if "약정서" in path.name or "체결" in path.name or "날인" in path.name:
        return 100
    if "리스크심의" in path.name:
        return 90
    if "teaser" in name or "v2.0" in name:
        return 80
    if "주간" in path.name:
        return 70
    if "정상화" in path.name or "rescue" in name or "f2" in name or "추진계획" in path.name:
        return 40
    return 50


def extract_text(path: Path) -> tuple[str, str]:
    suffix = path.suffix.lower()
    try:
        if suffix in {".xlsx", ".xlsm"}:
            wb = load_workbook(path, data_only=True, read_only=True)
            chunks = []
            for ws in wb.worksheets:
                chunks.append(f"[sheet:{ws.title}]")
                for row in ws.iter_rows(values_only=True):
                    values = [str(v) for v in row if v is not None]
                    if values:
                        chunks.append(" | ".join(values))
                    if len(chunks) > 2600:
                        break
            return "\n".join(chunks), "표/텍스트추출"
        if suffix == ".docx":
            import docx  # type: ignore

            doc = docx.Document(path)
            chunks = [p.text for p in doc.paragraphs if p.text]
            for table in doc.tables[:20]:
                for row in table.rows:
                    values = [c.text.strip() for c in row.cells if c.text.strip()]
                    if values:
                        chunks.append(" | ".join(values))
            return "\n".join(chunks), "표/텍스트추출"
        if suffix == ".pdf":
            try:
                from pypdf import PdfReader  # type: ignore
            except Exception:
                from PyPDF2 import PdfReader  # type: ignore

            reader = PdfReader(str(path))
            pages = []
            for page in reader.pages[:80]:
                try:
                    pages.append(page.extract_text() or "")
                except Exception:
                    continue
            return "\n".join(pages), "수동검증필요"
        if suffix == ".pptx":
            chunks = []
            with zipfile.ZipFile(path) as zf:
                for name in sorted(n for n in zf.namelist() if n.startswith("ppt/slides/slide") and n.endswith(".xml")):
                    root = ET.fromstring(zf.read(name))
                    texts = [node.text for node in root.iter() if node.tag.endswith("}t") and node.text]
                    if texts:
                        chunks.append(" ".join(texts))
            return "\n".join(chunks), "텍스트추출"
        if suffix == ".doc":
            return "", "수동검증필요"
    except Exception as exc:
        return f"[extract_error] {exc}", "수동검증필요"
    return "", "수동검증필요"


def build_source_index() -> tuple[list[dict[str, Any]], dict[str, int]]:
    files = [
        p
        for p in IOTA_DIR.rglob("*")
        if p.is_file()
        and "_v0.10" not in p.name
        and "v0.10_" not in p.name
        and not p.name.startswith("~$")
    ]
    hashes = {p: sha256(p) for p in files}
    hash_counts = Counter(hashes.values())
    rows = []
    for idx, path in enumerate(sorted(files, key=lambda p: str(p.relative_to(IOTA_DIR))), start=1):
        stat = path.stat()
        rel = str(path.relative_to(IOTA_DIR))
        rows.append(
            {
                "Source_ID": f"SRC010-{idx:04d}",
                "Path": rel,
                "File_Name": path.name,
                "Extension": path.suffix.lower(),
                "Last_Modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                "Size_Bytes": stat.st_size,
                "SHA256": hashes[path],
                "Duplicate_Count": hash_counts[hashes[path]],
                "Is_Duplicate": "Y" if hash_counts[hashes[path]] > 1 else "N",
                "Document_Date": infer_source_date(path.name, stat.st_mtime),
                "Document_Version": infer_version(path.name),
                "Priority": source_priority(path),
            }
        )
    return rows, dict(Counter(row["Extension"] for row in rows))


def numeric_candidates(source_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    out = []
    candidate_id = 1
    amount_re = re.compile(r"(?P<num>\d{1,3}(?:,\d{3})*(?:\.\d+)?|\d+(?:\.\d+)?)\s*(?P<unit>억원|백만원|만원|원|㎡|평|%|대|개월|년|일)")
    for source in source_rows:
        path = IOTA_DIR / source["Path"]
        text, quality = extract_text(path)
        if not text:
            continue
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        count_for_file = 0
        for line_no, line in enumerate(lines, start=1):
            for match in amount_re.finditer(line):
                value = match.group("num")
                unit = match.group("unit")
                category = classify_candidate(line, unit)
                out.append(
                    {
                        "Candidate_ID": f"NC010-{candidate_id:05d}",
                        "Source_ID": source["Source_ID"],
                        "File_Name": source["File_Name"],
                        "Line_or_Cell": line_no,
                        "Category": category,
                        "Raw_Value": value,
                        "Unit": unit,
                        "Context": line[:280],
                        "Source_Date": source["Document_Date"],
                        "Document_Version": source["Document_Version"],
                        "Extraction_Quality": quality,
                    }
                )
                candidate_id += 1
                count_for_file += 1
                if count_for_file >= 200:
                    break
            if count_for_file >= 200:
                break
    return out


def classify_candidate(context: str, unit: str) -> str:
    text = context.lower()
    if unit in {"억원", "백만원", "만원", "원"}:
        if any(k in context for k in ["공사비", "도급", "공사"]):
            return "공사비"
        if any(k in context for k in ["대출", "Loan", "loan", "Tr.", "트랜치", "금리"]):
            return "대출"
        if any(k in context for k in ["Equity", "투자", "주주", "출자", "자본"]):
            return "지분"
        return "금액"
    if unit in {"㎡", "평", "대"}:
        return "면적/물리제원"
    if unit == "%":
        if any(k in text for k in ["ltv", "dscr", "금리", "interest"]):
            return "대출"
        return "비율"
    return "일정"


def lfc_payload() -> dict[str, Any]:
    equity_427 = round(
        sum(row[3] for row in EQUITY_ROWS if row[0] == "427" and row[2] not in EQUITY_427_DETAIL_ROWS)
    )
    equity_816 = round(sum(row[3] for row in EQUITY_ROWS if row[0] == "816"))
    shareholder_427 = 2000
    shareholder_816 = 2535
    loan_427 = 48000
    loan_816 = sum(row[2] for row in REFI_816_ROWS)
    total_equity = equity_427 + equity_816
    total_shareholder = shareholder_427 + shareholder_816
    total_loan = loan_427 + loan_816
    total_aum = total_equity + total_shareholder + total_loan

    equity_investors = []
    for project, tranche, investor, amount, timing in EQUITY_ROWS:
        row: dict[str, Any] = {"project": project, "tranche": tranche, "investor": investor, "amount": amount, "timing": timing}
        if project == "427" and investor in EQUITY_427_DETAIL_ROWS:
            row["countInTotal"] = False
        equity_investors.append(row)
    loan_lenders = [
        {
            "id": "427-senior-frame",
            "project": "427",
            "loanType": "본PF / Senior Loan",
            "tranche": "",
            "lender": "",
            "amount": loan_427,
            "execution": "",
            "maturity": "",
            "rate": "",
            "fee": "",
        }
    ]
    for idx, (tranche, lender, amount) in enumerate(REFI_816_ROWS, start=1):
        loan_lenders.append(
            {
                "id": f"816-refi-{idx:02d}",
                "project": "816",
                "loanType": "Refi",
                "tranche": tranche,
                "lender": lender,
                "amount": amount,
                "execution": "2026-04-23",
                "maturity": "2027-04-23",
                "rate": "",
                "fee": "",
            }
        )

    lender_groups = lender_search_groups(loan_lenders)
    return {
        "meta": {"lead": "박준호", "members": "1명"},
        "summaryCards": [
            {"id": "aum", "label": "AUM", "caption": "Equity, 주주대여금, Loan 합산 관리 규모"},
            {"id": "equity", "label": "Equity", "caption": "Vehicle별 투자자 투입 자본"},
            {"id": "loan", "label": "Loan", "caption": "대주별 약정액과 tranche 관리"},
            {"id": "status", "label": "주요 현황", "caption": "LFC 관리 대상과 다음 데이터 확보 항목"},
        ],
        "cardDetails": {
            "aum": [
                ["Equity", equity_427, equity_816, total_equity, "최종 엑셀 수령 후 교체"],
                ["주주대여금", shareholder_427, shareholder_816, total_shareholder, "최종 엑셀 수령 후 교체"],
                ["Loan", loan_427, loan_816, total_loan, "최종 엑셀 수령 후 교체"],
                ["AUM", equity_427 + shareholder_427 + loan_427, equity_816 + shareholder_816 + loan_816, total_aum, "최종 엑셀 수령 후 교체"],
            ],
            "equity": [
                ["PFV 직접 Equity", equity_427, equity_816, total_equity, "최종 엑셀 수령 후 교체"],
                ["421호/REF 계열", 591, 20, 611, "최종 엑셀 수령 후 교체"],
                ["시공/전략 투자자", 99, 6, 105, "최종 엑셀 수령 후 교체"],
                ["기타 금융/고유 투자자", 110, 0, 110, "최종 엑셀 수령 후 교체"],
            ],
            "loan": [
                ["Senior / 본PF Loan", loan_427, "", loan_427, "최종 엑셀 수령 후 교체"],
                ["Refi Loan", "", loan_816, loan_816, "최종 엑셀 수령 후 교체"],
                ["현재 Loan 합계", loan_427, loan_816, total_loan, "최종 엑셀 수령 후 교체"],
                ["대체 전 B/L", "", 7170, 7170, "히스토리"],
            ],
            "status": [
                ["427", "본PF 개별 대주명·조건표 확보 필요", "", "", ""],
                ["816", "Refi 7,970억원 대주별 명단 반영", "", "", ""],
                ["통합", "427 + 816 기준으로 합산", "", "", ""],
                ["뉴스", "대주 테이블 기준 검색대상 자동 생성", "", "", ""],
            ],
        },
        "capitalStacks": {
            "427": {
                "label": "IOTA One 427",
                "segments": [
                    {"name": "Equity", "value": equity_427, "type": "equity"},
                    {"name": "주주대여금", "value": shareholder_427, "type": "shareholder"},
                    {"name": "Loan", "value": loan_427, "type": "loan"},
                ],
            },
            "816": {
                "label": "IOTA Two 816",
                "segments": [
                    {"name": "Equity", "value": equity_816, "type": "equity"},
                    {"name": "주주대여금", "value": shareholder_816, "type": "shareholder"},
                    {"name": "Loan", "value": loan_816, "type": "loan"},
                ],
            },
            "total": {
                "label": "통합",
                "segments": [
                    {"name": "Equity", "value": total_equity, "type": "equity"},
                    {"name": "주주대여금", "value": total_shareholder, "type": "shareholder"},
                    {"name": "Loan", "value": total_loan, "type": "loan"},
                ],
            },
        },
        "equityInvestors": equity_investors,
        "loanLenders": loan_lenders,
        "marketNewsFallback": {
            "generatedAt": datetime.now().astimezone().isoformat(timespec="seconds"),
            "windowDays": 3,
            "items": [
                {
                    "lender": g["lender"],
                    "relation": g["relation"],
                    "projects": g.get("projects", []),
                    "articles": [],
                }
                for g in lender_groups
            ],
        },
        "newsLenders": lender_groups,
        "extraItems": [
            ["대출 조건 체크리스트", "대출금액, 금리, 수수료, 선후순위, 인출조건, 후행조건을 조건별로 체크합니다."],
            ["만기 캘린더", "Vehicle별 만기, 연장 가능일, 본PF 전환 목표일을 캘린더형으로 관리합니다."],
            ["자료 요청 리스트", "대주별 약정서, term sheet, 수수료 산출표, 신용보강 자료 요청 상태를 관리합니다."],
            ["대주 커뮤니케이션 로그", "대주별 회신, 요청사항, 다음 미팅일, 내부 담당자를 누적 관리합니다."],
            ["본PF 전환 준비사항", "Refi 이후 본PF 전환을 위한 조건, 모집 주관사, 대체 상환재원을 별도 추적합니다."],
        ],
    }


def lender_search_groups(loan_lenders: list[dict[str, Any]]) -> list[dict[str, Any]]:
    groups: dict[str, dict[str, Any]] = {}
    for row in loan_lenders:
        lender = row.get("lender") or ""
        if not lender:
            continue
        search_terms = lender_terms(lender)
        display = search_terms[0] if search_terms else lender
        relation = f"{row.get('project')} {row.get('tranche') or row.get('loanType')}"
        if display not in groups:
            groups[display] = {"lender": display, "relation": relation, "searchNames": search_terms, "projects": [row.get("project")]}
        else:
            groups[display]["relation"] = f"{groups[display]['relation']} / {relation}"
            if row.get("project") not in groups[display]["projects"]:
                groups[display]["projects"].append(row.get("project"))
    defaults = [
        ("KB국민은행", "427 본PF 후보 주관기관 모니터링", ["KB국민은행", "KB금융"], ["427"]),
        ("신한GIB", "427 금융 투자자·본PF 후보군 모니터링", ["신한GIB", "신한금융"], ["427"]),
        ("신한은행", "427 금융 투자자·대주 후보군 모니터링", ["신한은행", "신한금융"], ["427"]),
        ("신한캐피탈", "427 금융 투자자·여전사 모니터링", ["신한캐피탈", "신한금융"], ["427"]),
        ("신한투자증권", "427/816 투자자·SPC 관련 모니터링", ["신한투자증권", "신한금융"], ["427", "816"]),
        ("NH투자증권", "816 Tr.A-2 SPC 및 대리금융 관련 모니터링", ["NH투자증권", "NH금융"], ["816"]),
    ]
    for display, relation, search_names, projects in defaults:
        if display not in groups:
            groups[display] = {"lender": display, "relation": relation, "searchNames": search_names, "projects": projects}
        else:
            for project in projects:
                if project not in groups[display]["projects"]:
                    groups[display]["projects"].append(project)
    return sorted(groups.values(), key=lambda x: x["lender"])


def lender_terms(lender: str) -> list[str]:
    mapping = [
        ("메리츠", "메리츠화재"),
        ("갤럭시이오", "NH투자증권"),
        ("816공간제일차", "신한투자증권"),
        ("한화실버아이언", "한화저축은행"),
        ("한국투자Debt", "한국투자리얼에셋운용"),
        ("한국투자메자닌", "한국투자리얼에셋운용"),
        ("스틱", "스틱얼터너티브자산운용"),
        ("대신저축은행", "대신저축은행"),
        ("비씨카드", "비씨카드"),
        ("흥국저축은행", "흥국저축은행"),
        ("키움", "키움투자자산운용"),
        ("이터널하이브", "대신증권"),
        ("코람코", "코람코자산운용"),
        ("케이에이치엘제이십일차", "메리츠증권"),
        ("케이에이치엘제이십이차", "소노인터내셔널"),
    ]
    for needle, term in mapping:
        if needle in lender:
            return [term, lender]
    return [lender]


def build_extra_rows(lfc: dict[str, Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    cid = 1
    refi_file = find_file("YD816PFV_bridge loan_대주리스트") or Path("YD816PFV_bridge loan_대주리스트.xlsx")
    equity_file = find_file("vehicle별 equity") or Path("vehicle별 equity 투자자 현황.xlsx")

    for project, tranche, investor, amount, timing in EQUITY_ROWS:
        rows.append(
            extra_row(
                cid,
                equity_file,
                "금융",
                "Equity",
                project,
                f"{project} {investor} Equity",
                amount,
                "억원",
                "10_TB_COMPANY / 30_TB_CAPITAL_STACK",
                "LFC Equity 투자자 현황",
            )
        )
        cid += 1
    for tranche, lender, amount in REFI_816_ROWS:
        rows.append(
            extra_row(
                cid,
                refi_file,
                "금융",
                "Loan",
                "816 Refi",
                f"{tranche} {lender}",
                amount,
                "억원",
                "30_TB_CAPITAL_STACK",
                "LFC Loan 대주 현황",
            )
        )
        cid += 1
    for key, stack in lfc["capitalStacks"].items():
        for seg in stack["segments"]:
            rows.append(
                extra_row(
                    cid,
                    refi_file if seg["name"] == "Loan" and key == "816" else equity_file,
                    "금융",
                    "Capital Stack",
                    stack["label"],
                    f"{stack['label']} {seg['name']}",
                    seg["value"],
                    "억원",
                    "30_TB_CAPITAL_STACK",
                    "LFC Funding Structure",
                )
            )
            cid += 1
    return rows


def extra_row(
    cid: int,
    source_path: Path,
    big: str,
    mid: str,
    small: str,
    item: str,
    value: Any,
    unit: str,
    target_sheet: str,
    screen: str,
) -> dict[str, Any]:
    source_name = source_path.name
    return {
        "Candidate_ID": f"EX010-{cid:04d}",
        "Source_ID": slug(source_name),
        "Source_Date": infer_source_date(source_name, source_path.stat().st_mtime if source_path.exists() else None),
        "Document_Version": infer_version(source_name),
        "대분류": big,
        "중분류": mid,
        "소분류": small,
        "항목": item,
        "값": value,
        "단위": unit,
        "근거": source_name,
        "원문 존재 확실성": "원문 확인",
        "사업 확정 상태": "최종 엑셀 대기",
        "Target_Sheet": target_sheet,
        "Target_PK": "",
        "Scenario_Tag": "",
        "Value_Status": "최종파일대기",
        "Promotion_Status": "화면프레임",
        "Conflict_Group": "",
        "Superseded_By": "",
        "Extraction_Quality": "표/텍스트추출",
        "화면추천위치": screen,
    }


def slug(text: str) -> str:
    return re.sub(r"[^0-9A-Za-z가-힣]+", "_", text).strip("_")[:80] or "unknown"


def style_sheet(ws) -> None:
    fill = PatternFill("solid", fgColor="1F2A44")
    font = Font(color="FFFFFF", bold=True)
    for cell in ws[1]:
        cell.fill = fill
        cell.font = font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions
    for column in ws.columns:
        letter = get_column_letter(column[0].column)
        width = 12
        for cell in column[:250]:
            if cell.value is not None:
                width = max(width, min(52, len(str(cell.value)) + 2))
        ws.column_dimensions[letter].width = width
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)


def replace_sheet(wb, name: str, headers: list[str], rows: list[list[Any]]) -> None:
    if name in wb.sheetnames:
        del wb[name]
    ws = wb.create_sheet(name)
    ws.append(headers)
    for row in rows:
        ws.append(row)
    style_sheet(ws)


def add_v010_sheets(workbook_path: Path, source_rows: list[dict[str, Any]], numeric_rows: list[dict[str, Any]], extra_rows: list[dict[str, Any]], lfc: dict[str, Any]) -> None:
    wb = load_workbook(workbook_path)
    for old in ["98_EXTRA_SOURCE_EXTRACTS_v0.9", "98_EXTRA_SOURCE_EXTRACTS_v0.8"]:
        if old in wb.sheetnames:
            del wb[old]
    replace_sheet(wb, EXTRA_SHEET, EXTRA_HEADERS, [[r.get(h, "") for h in EXTRA_HEADERS] for r in extra_rows])
    replace_sheet(wb, SOURCE_INDEX_SHEET, list(source_rows[0].keys()), [[r.get(h, "") for h in source_rows[0].keys()] for r in source_rows])
    replace_sheet(wb, NUMERIC_LOG_SHEET, list(numeric_rows[0].keys()) if numeric_rows else ["Candidate_ID"], [[r.get(h, "") for h in numeric_rows[0].keys()] for r in numeric_rows] if numeric_rows else [])
    latest_headers = ["Decision_ID", "Item", "Selected_Value", "Selected_Source", "Rule", "Decision", "Note"]
    latest_rows = [
        ["LD010-0001", "816 Refi 대주별 약정액", "최종 엑셀 대기", "YD816PFV_bridge loan_대주리스트.xlsx", "사용자 지시에 따라 금액 검증 보류", "보류", "화면 프레임용 행만 유지"],
        ["LD010-0002", "427 Equity", "최종 엑셀 대기", "vehicle별 equity 투자자 현황.xlsx", "사용자 지시에 따라 금액 검증 보류", "보류", "화면 프레임용 행만 유지"],
        ["LD010-0003", "816 Equity", "최종 엑셀 대기", "vehicle별 equity 투자자 현황.xlsx", "사용자 지시에 따라 금액 검증 보류", "보류", "화면 프레임용 행만 유지"],
        ["LD010-0004", "427 개별 대주명", "", "iota one 427 pre pf 대출 약정서.DOC", ".DOC 수동검증 필요", "보류", "임의 분해하지 않음"],
    ]
    replace_sheet(wb, LATESTNESS_LOG_SHEET, latest_headers, latest_rows)
    lfc_headers = ["Type", "Project", "Loan_Type", "Tranche", "Name", "Amount_억원", "Execution", "Maturity", "Rate", "Fee"]
    lfc_rows = []
    for row in lfc["equityInvestors"]:
        lfc_rows.append(["Equity", row["project"], "", row["tranche"], row["investor"], row["amount"], row.get("timing", ""), "", "", ""])
    for row in lfc["loanLenders"]:
        lfc_rows.append(["Loan", row["project"], row["loanType"], row["tranche"], row["lender"], row["amount"], row["execution"], row["maturity"], row["rate"], row["fee"]])
    replace_sheet(wb, LFC_DATA_SHEET, lfc_headers, lfc_rows)
    wb.save(workbook_path)


def write_handoff(source_rows: list[dict[str, Any]], extra_rows: list[dict[str, Any]], numeric_rows: list[dict[str, Any]]) -> None:
    wb = Workbook()
    wb.remove(wb.active)

    def add(name: str, headers: list[str], rows: list[list[Any]]) -> None:
        ws = wb.create_sheet(name)
        ws.append(headers)
        for row in rows:
            ws.append(row)
        style_sheet(ws)

    unfilled_rows = read_unfilled(MASTER_V010)
    confirmed = [r for r in extra_rows if r["Promotion_Status"] == "반영"]
    waiting = [r for r in extra_rows if r["Promotion_Status"] == "화면프레임"]
    summary = [
        ["버전", "v0.10"],
        ["기준 입력본", "IOTA_Seoul_Master_DB_v0.9.xlsx / IOTA_Seoul_DB_for_UI_v0.9.xlsx"],
        ["전체 원천 파일", len(source_rows)],
        ["원천추출 후보", len(extra_rows)],
        ["숫자 후보", len(numeric_rows)],
        ["확정값", len(confirmed)],
        ["최종 엑셀 대기 프레임값", len(waiting)],
        ["미충전값", len(unfilled_rows)],
        ["LFC 보강", "화면 프레임과 대주/투자자 행 구조 생성. Equity/Loan 금액 검증은 최종 엑셀 수령 후 진행"],
    ]
    add("요약", ["구분", "값"], summary)

    def handoff_row(r: dict[str, Any]) -> list[Any]:
        return [r["대분류"], r["중분류"], r["소분류"], r["항목"], r["값"], r["단위"], r["사업 확정 상태"], r.get("화면추천위치", "LFC")]

    add("확정값", HANDOFF_HEADERS, [handoff_row(r) for r in confirmed])
    add("계획_시나리오값", HANDOFF_HEADERS, [handoff_row(r) for r in waiting])
    add("충돌값", HANDOFF_HEADERS, [])
    add("미충전값", ["시트명", "행 식별자", "미충전 컬럼", "사유", "필요자료"], unfilled_rows)
    add("원천추출", EXTRA_HEADERS, [[r.get(h, "") for h in EXTRA_HEADERS] for r in extra_rows])
    add("파일목록", list(source_rows[0].keys()), [[r.get(h, "") for h in source_rows[0].keys()] for r in source_rows])
    add("숫자후보", list(numeric_rows[0].keys()) if numeric_rows else ["Candidate_ID"], [[r.get(h, "") for h in numeric_rows[0].keys()] for r in numeric_rows] if numeric_rows else [])
    wb.save(HANDOFF_V010)


def read_unfilled(path: Path) -> list[list[Any]]:
    if not path.exists():
        return []
    wb = load_workbook(path, data_only=True, read_only=True)
    if UNFILLED_SHEET not in wb.sheetnames:
        return []
    ws = wb[UNFILLED_SHEET]
    headers = [c.value for c in ws[1]]
    rows = []
    for raw in ws.iter_rows(min_row=2, values_only=True):
        if not any(raw):
            continue
        d = dict(zip(headers, raw))
        rows.append(
            [
                d.get("시트명") or d.get("Sheet"),
                d.get("행 식별자 (PK)") or d.get("PK"),
                d.get("미충전 컬럼명") or d.get("Column"),
                d.get("사유") or d.get("Reason"),
                d.get("채우려면 필요한 자료 (예상 파일명 또는 유형)") or d.get("Needed"),
            ]
        )
    return rows


def workbook_summary(path: Path) -> dict[str, Any]:
    wb = load_workbook(path, data_only=True, read_only=True)
    return {
        "file": path.name,
        "sheets": len(wb.sheetnames),
        "unfilled_rows": max(0, wb[UNFILLED_SHEET].max_row - 1) if UNFILLED_SHEET in wb.sheetnames else 0,
        "source_index_rows": max(0, wb[SOURCE_INDEX_SHEET].max_row - 1) if SOURCE_INDEX_SHEET in wb.sheetnames else 0,
        "extra_rows": max(0, wb[EXTRA_SHEET].max_row - 1) if EXTRA_SHEET in wb.sheetnames else 0,
    }


def inject_html(lfc: dict[str, Any], source_count: int, extra_count: int) -> None:
    if not PUBLIC_HTML.exists():
        return
    html = PUBLIC_HTML.read_text(encoding="utf-8")
    payload = {
        "version": "v0.10",
        "workbook": "IOTA_Seoul_DB_for_UI_v0.10.xlsx",
        "generatedAt": datetime.now().strftime("%Y-%m-%d"),
        "sourceDate": "2026-04-30",
        "sourceFileCount": source_count,
        "extraExtractRows": extra_count,
        "financeValues": {
            "Equity/Loan": "최종 엑셀 수령 후 금액 검증 및 교체 예정",
        },
        "syncNotes": [
            "HTML LFC 표시값은 v0.10 UI workbook 기준으로 동기화합니다.",
            "Equity/Loan 금액은 사용자 최종 엑셀 수령 전까지 검증하지 않습니다.",
            "원천에서 확인되지 않은 427 개별 대주명은 빈칸으로 유지합니다.",
        ],
        "lfc": lfc,
    }
    block = (
        "    // @iota-data-v010:start\n"
        f"    const iotaData = {json.dumps(payload, ensure_ascii=False, indent=6)};\n"
        "    // @iota-data-v010:end"
    )
    html = re.sub(
        r"    // @iota-data-v0?9:start[\s\S]*?// @iota-data-v0?9:end|    // @iota-data-v010:start[\s\S]*?// @iota-data-v010:end",
        block,
        html,
        count=1,
    )
    html = re.sub(
        r"    const LFC_V2 = \{[\s\S]*?\n    \};\n\n    const app =",
        "    const LFC_V2 = iotaData.lfc;\n\n    const app =",
        html,
        count=1,
    )
    PUBLIC_HTML.write_text(html, encoding="utf-8")
    DOCS_HTML.write_text(html, encoding="utf-8")


def write_lender_json(lfc: dict[str, Any]) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    lender_payload = {"generatedAt": datetime.now().isoformat(timespec="seconds"), "lenders": lfc["newsLenders"]}
    (DATA_DIR / "lfc-lenders.json").write_text(json.dumps(lender_payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_package_index() -> None:
    PACKAGE_DIR.mkdir(parents=True, exist_ok=True)
    readme = """# IOTA Seoul Data Package v0.10

이 폴더는 IOTA Seoul 개발관리 Workspace와 LFC 화면 제작을 위한 v0.10 데이터 패키지입니다.
v0.10은 기존 v0.9를 덮어쓰지 않고, `iota_reference` 전체 82개 파일을 다시 색인한 뒤 LFC 화면 프레임을 보강한 버전입니다. Equity/Loan 금액은 최종 엑셀 수령 전까지 확정값으로 검증하지 않습니다.

## 파일 구성

| 파일명 | 목적 | 보는 방법 |
|---|---|---|
| `IOTA_Seoul_Master_DB_v0.10.xlsx` | 원본 스키마를 유지한 마스터 DB입니다. v0.10 원천 파일 색인, 숫자 후보, 최신성 판단 로그, LFC 데이터 시트가 추가되어 있습니다. | 데이터 검수자나 PM이 전체 DB 구조와 원천 반영 상태를 볼 때 사용합니다. |
| `IOTA_Seoul_DB_for_UI_v0.10.xlsx` | UI/UX 제작자가 화면에 바로 쓰기 쉽도록 정리한 버전입니다. | 화면 제작자는 이 파일을 우선 참고합니다. |
| `IOTA_Seoul_IFPDP_HandOff_v0.10.xlsx` | 제작 담당자 전달용 요약 데이터입니다. `요약 → 확정값 → 계획_시나리오값 → 충돌값 → 미충전값 → 원천추출 → 파일목록` 순서로 구성했습니다. | 화면 구성이나 디자인 시 필요한 값만 빠르게 찾을 때 사용합니다. |
| `v0.10_build_report.md` | v0.10 생성 결과, 원천 파일 수, 숫자 후보 수, LFC 보강 내역, 구조 검증 결과를 정리한 보고서입니다. | 데이터 품질과 남은 확인사항을 볼 때 사용합니다. |
| `IOTA_Seoul_Data_Package_v0.10.zip` | 위 파일과 이 안내문을 묶은 다운로드용 압축파일입니다. | 팀원에게 한 번에 전달할 때 사용합니다. |

## v0.10에서 특히 볼 부분

1. `IOTA_Seoul_IFPDP_HandOff_v0.10.xlsx`의 `요약` 시트를 먼저 봅니다.
2. LFC 화면 제작값은 `IOTA_Seoul_DB_for_UI_v0.10.xlsx`의 `v0.10_LFC_DATA`를 봅니다.
3. LFC 투자자/대주 테이블은 화면 제작용 행 구조를 우선 구성했습니다.
4. Equity/Loan 금액은 사용자가 제공할 최종 엑셀 파일을 받은 뒤 확정 검증 및 교체합니다.
5. 원천에서 확인되지 않은 427 개별 대주명, 금리, 수수료는 빈칸으로 유지했습니다.

## 주의사항

- 원천에서 확인되지 않은 값은 임의로 채우지 않았습니다.
- `.DOC` 파일은 안정적인 자동 추출이 어려워 수동검증필요로 남겼습니다.
- Rescue/F2/정상화계획안/PF 추진계획안의 값은 확정값과 구분해서 봐야 합니다.
- 민감한 내부 협상 전략, 금리·수수료·특약, 임차 조건, 내부 반대 의견 등은 외부 제출용 화면에서 제외하거나 마스킹해야 합니다.
"""
    (PACKAGE_DIR / "README.md").write_text(readme, encoding="utf-8")
    index = """<!doctype html>
<html lang="ko">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>IOTA Seoul Data Package v0.10</title>
  <style>
    :root { color-scheme: light; --text:#1d1d1f; --muted:#6e6e73; --line:#d8d8dc; --panel:#fff; --bg:#f5f5f7; --accent:#1f2a44; }
    * { box-sizing: border-box; }
    body { margin:0; font-family:-apple-system,BlinkMacSystemFont,"Segoe UI","Noto Sans KR",Arial,sans-serif; background:var(--bg); color:var(--text); line-height:1.6; }
    main { max-width:960px; margin:0 auto; padding:56px 24px 72px; }
    h1 { margin:0 0 10px; font-size:34px; line-height:1.15; }
    .lead { margin:0 0 28px; color:var(--muted); font-size:16px; }
    .card { background:var(--panel); border:1px solid var(--line); border-radius:14px; padding:22px; margin-top:16px; }
    .grid { display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:14px; }
    a { color:var(--accent); font-weight:800; text-decoration:none; }
    a:hover { text-decoration:underline; }
    .file { display:block; border:1px solid var(--line); border-radius:12px; padding:14px; min-height:106px; background:#fbfbfd; }
    .file strong { display:block; margin-bottom:6px; font-size:15px; }
    .file span { display:block; color:var(--muted); font-size:13px; }
    ol { margin:0; padding-left:21px; }
    @media (max-width:720px){ .grid{grid-template-columns:1fr;} main{padding:36px 16px 54px;} }
  </style>
</head>
<body>
  <main>
    <h1>IOTA Seoul Data Package v0.10</h1>
    <p class="lead">IOTA Seoul 개발관리 Workspace 및 LFC 화면 제작/데이터 검수용 산출물입니다. 기준일은 2026-04-30입니다.</p>
    <section class="card">
      <h2>바로 다운로드</h2>
      <div class="grid">
        <a class="file" href="./IOTA_Seoul_Data_Package_v0.10.zip" download><strong>전체 패키지 ZIP</strong><span>엑셀 3종, build report, README.md를 한 번에 받습니다.</span></a>
        <a class="file" href="./README.md"><strong>README.md</strong><span>각 파일의 목적과 확인 순서를 설명한 안내문입니다.</span></a>
        <a class="file" href="./IOTA_Seoul_IFPDP_HandOff_v0.10.xlsx" download><strong>IFPDP HandOff v0.10</strong><span>제작 담당자가 먼저 볼 요약형 엑셀입니다.</span></a>
        <a class="file" href="./IOTA_Seoul_DB_for_UI_v0.10.xlsx" download><strong>DB for UI v0.10</strong><span>UI 화면 제작에 바로 쓰기 위한 정제 DB입니다.</span></a>
        <a class="file" href="./IOTA_Seoul_Master_DB_v0.10.xlsx" download><strong>Master DB v0.10</strong><span>스키마와 검수용 메타를 유지한 마스터 DB입니다.</span></a>
        <a class="file" href="./v0.10_build_report.md"><strong>Build Report</strong><span>원천 파일 색인, LFC 보강, 구조 검증 결과입니다.</span></a>
      </div>
    </section>
    <section class="card">
      <h2>권장 확인 순서</h2>
      <ol>
        <li>먼저 <strong>IFPDP HandOff v0.10</strong>의 요약 시트를 봅니다.</li>
        <li>LFC 화면 제작값은 <strong>v0.10_LFC_DATA</strong>와 <strong>계획_시나리오값</strong>을 봅니다.</li>
        <li>충돌값과 미충전값은 추가 원천 확인 전까지 화면 확정값으로 쓰지 않습니다.</li>
      </ol>
    </section>
    <section class="card"><h2>관련 화면</h2><p><a href="../iota-development-workspace.html#lfc">IOTA LFC Workspace HTML 열기</a></p></section>
  </main>
</body>
</html>
"""
    (PACKAGE_DIR / "index.html").write_text(index, encoding="utf-8")


def copy_package_files() -> None:
    PACKAGE_DIR.mkdir(parents=True, exist_ok=True)
    for path in [MASTER_V010, UI_V010, HANDOFF_V010, REPORT_V010]:
        shutil.copy2(path, PACKAGE_DIR / path.name)
    zip_path = PACKAGE_DIR / "IOTA_Seoul_Data_Package_v0.10.zip"
    if zip_path.exists():
        zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for name in [
            "README.md",
            MASTER_V010.name,
            UI_V010.name,
            HANDOFF_V010.name,
            REPORT_V010.name,
        ]:
            zf.write(PACKAGE_DIR / name, arcname=name)


def write_report(source_rows: list[dict[str, Any]], numeric_rows: list[dict[str, Any]], extra_rows: list[dict[str, Any]], validations: list[dict[str, Any]], lfc: dict[str, Any]) -> None:
    ext_counts = Counter(r["Extension"] for r in source_rows)
    duplicate_count = sum(1 for r in source_rows if r["Is_Duplicate"] == "Y")
    lines = [
        "# IOTA Seoul v0.10 Build Report",
        "",
        f"- 생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"- 입력본: `{MASTER_V09.name}`, `{UI_V09.name}`, `{HANDOFF_V09.name}`",
        f"- 산출물: `{MASTER_V010.name}`, `{UI_V010.name}`, `{HANDOFF_V010.name}`",
        "",
        "## 원천 파일 재색인",
        "",
        f"- 전체 파일 수: {len(source_rows)}건",
        f"- 확장자별 수: {dict(ext_counts)}",
        f"- 중복 해시 파일 수: {duplicate_count}건",
        "",
        "## v0.10 보강 내역",
        "",
        f"- LFC 원천추출 후보: {len(extra_rows)}건",
        f"- 숫자 후보 로그: {len(numeric_rows)}건",
        "- Equity/Loan 금액 검증: 사용자 지시에 따라 보류",
        "- 427/816 투자자·대주 행 구조와 화면 프레임만 유지",
        "- 427 개별 대주명·금리·수수료는 최종 엑셀 수령 전까지 빈칸 유지",
        "",
        "## 검증 결과",
        "",
    ]
    for item in validations:
        lines.append(
            f"- {item['file']}: sheets={item['sheets']}, source_index_rows={item['source_index_rows']}, "
            f"extra_rows={item['extra_rows']}, unfilled_rows={item['unfilled_rows']}"
        )
    lines.extend(
        [
            "",
            "## 최신값 판단",
            "",
            "- Equity/Loan 관련 금액은 최종 엑셀 수령 전까지 확정값으로 검증하지 않습니다.",
            "- 현재 HTML과 HandOff에는 화면 제작을 위한 행 구조와 임시 프레임만 유지합니다.",
            "- `.DOC` 파일은 안정 자동 추출이 어려워 수동검증필요로 남겼고, 해당 파일의 값은 임의 승격하지 않았습니다.",
        ]
    )
    REPORT_V010.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    if not MASTER_V09.exists() or not UI_V09.exists() or not HANDOFF_V09.exists():
        raise RuntimeError("v0.9 입력 파일이 필요합니다.")
    shutil.copy2(MASTER_V09, MASTER_V010)
    shutil.copy2(UI_V09, UI_V010)
    source_rows, _ = build_source_index()
    numeric_rows = numeric_candidates(source_rows)
    lfc = lfc_payload()
    extra_rows = build_extra_rows(lfc)
    add_v010_sheets(MASTER_V010, source_rows, numeric_rows, extra_rows, lfc)
    add_v010_sheets(UI_V010, source_rows, numeric_rows, extra_rows, lfc)
    write_handoff(source_rows, extra_rows, numeric_rows)
    validations = [workbook_summary(MASTER_V010), workbook_summary(UI_V010), workbook_summary(HANDOFF_V010)]
    write_report(source_rows, numeric_rows, extra_rows, validations, lfc)
    inject_html(lfc, len(source_rows), len(extra_rows))
    write_lender_json(lfc)
    write_package_index()
    copy_package_files()
    print(json.dumps({"source_files": len(source_rows), "numeric_candidates": len(numeric_rows), "lfc_frame_rows": len(extra_rows)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
